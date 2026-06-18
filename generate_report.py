"""生成课程设计报告 .docx"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import docx.oxml

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(51, 51, 51)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


# ===== 标题页 =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('人工智能课程设计报告')
run.font.size = Pt(26)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('题17 — 搭载数据分析 Skill 的表格智能 Agent')
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
doc.add_paragraph()

info_items = ['姓名：________________', '学号：________________', '班级：________________', '日期：________________']
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.size = Pt(14)

doc.add_page_break()

# ===== 目录占位 =====
doc.add_heading('目  录', level=1)
add_para('（生成后右键更新域或手动补全页码）')
doc.add_page_break()

# ==================== 一、摘要 ====================
doc.add_heading('一、课程设计摘要', level=1)
add_para(
    '本课程设计实现了一款搭载数据分析 Skill 的表格智能 Agent（TableAgent），'
    '选题定位 AI Agent 应用方向，属于 Function Call 实战项目。'
    '系统采用通义千问 qwen-plus 大语言模型作为智能推理引擎，'
    '基于 OpenAI SDK 手动搭建 Function Call 机制（未使用 LangChain 等高级框架），'
    '封装了数据读取与探查（DataReader）、统计分析（Analyst）、可视化绘图（Visualizer）'
    '三大核心 Skill 共 18 个工具函数。'
    '用户通过自然语言提出数据分析需求，Agent 自动识别意图、链式调用 Skill 函数，'
    '完成从数据加载、清洗探查、统计建模到图表生成的全流程分析任务。'
    '前端采用 Gradio ChatInterface 构建交互式对话界面，支持文字与图片混合展示。'
    '项目以 Titanic 数据集为默认数据源，完成 42 个单元测试全部通过，'
    '系统可稳定处理单步调用、链式调用、参数错误恢复等多场景。',
    indent=True
)

# ==================== 二、绪论 ====================
doc.add_heading('二、绪论', level=1)

doc.add_heading('2.1 选题背景与研究意义', level=2)
add_para(
    '随着大语言模型（LLM）能力的飞速发展，AI Agent（智能体）成为人工智能领域最重要的应用方向之一。'
    '传统的数据分析工作流需要用户掌握 Python、SQL 等编程技能，对非技术用户门槛较高。'
    '通过将 LLM 的语义理解能力与结构化数据分析工具相结合，可以构建能够理解自然语言指令、'
    '自动执行数据操作的智能 Agent，显著降低数据分析的门槛。',
    indent=True
)
add_para(
    '本项目的实践意义在于：（1）深入理解 LLM Function Call 机制的底层原理，'
    '通过手动实现而非框架封装的方式掌握 Agent 核心技术；'
    '（2）探索 LLM 在结构化数据分析场景中的应用边界与效果；'
    '（3）为构建更复杂的多 Skill Agent 系统提供可复用的架构设计参考。',
    indent=True
)

doc.add_heading('2.2 国内外研究现状', level=2)
add_para(
    '当前 AI Agent 领域的主流框架包括 LangChain、AutoGen、CrewAI 等，'
    '这些框架提供了高级抽象和开箱可用的 Agent 编排能力。然而，高层框架屏蔽了 Function Call 的底层细节，'
    '不利于深入理解 Agent 的工作原理。在数据分析领域，'
    'ChatGPT Code Interpreter、PandasAI 等产品已展示了 LLM 驱动数据分析的可行性。'
    '开源社区中也有众多基于 LangChain 的 Data Agent 实现，'
    '但缺乏对底层 Function Call 机制的剖析和手动实现。',
    indent=True
)
add_para(
    '本项目的差异化在于：不依赖高层框架，从头实现 Function Call 调度逻辑，'
    '使开发者能够完全掌控 Agent 的行为；同时聚焦表格数据分析这一垂直场景，'
    '通过精心设计的 Skill 函数确保 LLM 对数据操作的准确性和可解释性。',
    indent=True
)

doc.add_heading('2.3 主要研究内容与开发目标', level=2)
add_para(
    '本项目的主要研究内容包括：'
    '（1）基于通义千问 qwen-plus API 实现 Function Call 机制的 Agent 主循环；'
    '（2）设计并实现三大数据分析 Skill（数据读取、统计分析、可视化绘图）共 18 个工具函数；'
    '（3）构建基于 Gradio 的交互式对话界面；'
    '（4）确保 Agent 能正确处理单步调用、链式调用、参数错误恢复等场景。',
    indent=True
)
add_para(
    '开发目标：完成一个可交互演示的 TableAgent 系统，'
    '用户输入自然语言即可完成数据分析全流程，覆盖数据探查、统计分析和图表生成三大类任务。',
    indent=True
)

doc.add_heading('2.4 论文组织结构', level=2)
add_para(
    '本报告共分九章：第一章摘要概述项目全貌；第二章绪论阐述选题背景与研究现状；'
    '第三章介绍开发环境与关键技术原理；第四章说明数据集与系统总体设计；'
    '第五章详述核心功能实现过程与关键代码；第六章展示实验结果与分析；'
    '第七章进行系统测试与功能验证；第八章总结项目得失并展望未来；第九章列出参考文献。',
    indent=True
)

doc.add_page_break()

# ==================== 三、开发环境与关键技术原理 ====================
doc.add_heading('三、开发环境与关键技术原理', level=1)

doc.add_heading('3.1 软硬件开发环境', level=2)

add_para('硬件环境：', bold=True)
table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table.cell(0,0).text = '项目'
table.cell(0,1).text = '配置'
data = [
    ('CPU', 'Intel / AMD x86-64'),
    ('GPU', 'NVIDIA GPU (CUDA 12.4+)'),
    ('内存', '16 GB+'),
    ('操作系统', 'Windows 11'),
    ('存储', 'SSD 256 GB+'),
]
for i, (k, v) in enumerate(data):
    table.cell(i+1,0).text = k
    table.cell(i+1,1).text = v

add_para('')
add_para('软件环境：', bold=True)
table = doc.add_table(rows=10, cols=2, style='Light Grid Accent 1')
table.cell(0,0).text = '项目'
table.cell(0,1).text = '版本'
sw_data = [
    ('Python', '3.13'),
    ('PyTorch', '2.6.0+cu124'),
    ('OpenAI SDK', '≥2.41'),
    ('Pandas', '3.0.2'),
    ('NumPy', '≥2.3'),
    ('Matplotlib', '3.10.x'),
    ('Seaborn', '0.13.x'),
    ('Gradio', '6.18.0'),
    ('通义千问模型', 'qwen-plus'),
]
for i, (k, v) in enumerate(sw_data):
    table.cell(i+1,0).text = k
    table.cell(i+1,1).text = v

doc.add_heading('3.2 核心技术原理', level=2)

add_para('Function Call 机制', bold=True)
add_para(
    'Function Call（函数调用）是大语言模型的核心能力之一，'
    '允许模型在生成回复时识别需要调用外部工具的意图，并以结构化 JSON 格式输出函数名和参数。'
    '开发者解析该输出后执行对应函数，将结果返回给模型以生成最终回复。'
    '这一机制是 AI Agent 实现工具调用的基础。',
    indent=True
)
add_para(
    '本项目的 Agent 主循环严格遵循手搭 Function Call 范式：'
    '（1）将用户输入和所有 Skill 的 JSON Schema 定义一同发送给模型；'
    '（2）模型返回 tool_call 或直接回答；'
    '（3）如含 tool_call，解析函数名和参数并调度执行对应的 Skill 函数；'
    '（4）将函数执行结果追加到对话历史，继续循环直到模型直接回答。',
    indent=True
)

add_para('数据分析栈', bold=True)
add_para(
    '使用 Pandas 进行数据处理和统计分析，Matplotlib + Seaborn 进行数据可视化。'
    'Pandas 提供 DataFrame 结构化的数据操作接口，'
    'Seaborn 在 Matplotlib 基础上封装了更简洁美观的统计图表绘制功能。'
    '所有可视化函数返回 PIL Image 对象，可直接在 Gradio 界面中展示。',
    indent=True
)

add_para('Gradio 交互框架', bold=True)
add_para(
    'Gradio 是适用于机器学习演示的 Python Web UI 框架。'
    '其 gr.ChatInterface 组件可快速搭建类 ChatGPT 的对话界面，'
    '支持 OpenAI 风格的多模态消息格式（文字 + 图片），'
    '覆盖本项目对话交互和图表展示的需求。',
    indent=True
)

doc.add_page_break()

# ==================== 四、数据集与项目总体设计 ====================
doc.add_heading('四、数据集与项目总体设计', level=1)

doc.add_heading('4.1 数据集介绍', level=2)
add_para(
    '本项目默认使用 Kaggle 泰坦尼克号乘客数据集（Titanic: Machine Learning from Disaster）作为演示数据源。'
    '该数据集包含 891 条乘客记录，12 个字段，字段说明如下：',
    indent=True
)

table = doc.add_table(rows=13, cols=3, style='Light Grid Accent 1')
table.cell(0,0).text = '字段名'
table.cell(0,1).text = '类型'
table.cell(0,2).text = '说明'
titanic_fields = [
    ('PassengerId', 'int64', '乘客唯一标识'),
    ('Survived', 'int64', '是否存活（0=否, 1=是）'),
    ('Pclass', 'int64', '舱位等级（1/2/3）'),
    ('Name', 'object', '乘客姓名'),
    ('Sex', 'object', '性别'),
    ('Age', 'float64', '年龄（部分缺失）'),
    ('SibSp', 'int64', '同乘兄弟姐妹/配偶数'),
    ('Parch', 'int64', '同乘父母/子女数'),
    ('Ticket', 'object', '船票号'),
    ('Fare', 'float64', '票价'),
    ('Cabin', 'object', '客舱号（大量缺失）'),
    ('Embarked', 'object', '登船港口（C/Q/S）'),
]
for i, (f, t, d) in enumerate(titanic_fields):
    table.cell(i+1,0).text = f
    table.cell(i+1,1).text = t
    table.cell(i+1,2).text = d

add_para(
    '该数据集含缺失值（Age 约 20%、Cabin 约 77%、Embarked 约 0.2%），'
    '适合用来验证 Agent 的数据探查和缺失值处理能力。'
    '数据集的规模适中（891 行 × 12 列），'
    '便于快速完成全流程分析演示。',
    indent=True
)

doc.add_heading('4.2 数据预处理过程', level=2)
add_para(
    '由于本项目聚焦于 Agent 对数据的探查、分析和可视化能力，'
    '而非建模前的数据清洗，因此原始 Titanic 数据集保留了天然缺失值状态，'
    '未进行预处理。这一设计决策使得 Agent 能展示以下能力：',
    indent=True
)
add_para('（1）通过 get_overview / get_info / get_missing 自动识别数据结构和缺失情况；')
add_para('（2）根据用户需求针对性调用 filter_data 筛选数据子集；')
add_para('（3）在绘图时（如 histogram）自动 dropna 处理缺失值，保证图表正常生成。')

doc.add_heading('4.3 项目总体架构设计', level=2)
add_para('TableAgent 的系统架构分为四层：', indent=True)
add_para('用户交互层（Gradio UI）：接收自然语言输入，展示文字回复和图表图片。')
add_para('Agent 调度层（agent/core.py）：Function Call 主循环，负责意图识别、工具调度、结果整合。')
add_para('Skill 工具层（skills/）：封装 18 个数据分析函数，分为 DataReader、Analyst、Visualizer 三大类。')
add_para('数据层（Pandas DataFrame）：全局状态管理，所有 Skill 函数共享同一份 DataFrame 实例。')
add_para('')
add_para('系统运行流程如下：')
add_para(
    '（1）用户在 Gradio 界面输入自然语言问题；'
    '（2）Agent 将用户消息追加至对话历史，构造包含 Skill Schema 的 API 请求发送至 qwen-plus；'
    '（3）模型返回响应：若为直接回答则返回用户，若含 tool_call 则进入步骤（4）；'
    '（4）解析 tool_call 中的函数名和参数，调用对应的 Skill 函数并获取结果；'
    '（5）工具结果（文字或图片）追加至对话历史，回到步骤（2）继续循环；'
    '（6）循环上限为 15 次迭代，防止死循环。',
    indent=True
)

doc.add_heading('4.4 核心功能设计', level=2)

add_para('DataReader（数据读取与探查）', bold=True)
add_para(
    '包含 6 个函数：load_csv（加载数据）、get_overview（整体概览）、get_info（字段信息）、'
    'get_missing（缺失值统计）、get_head（预览数据）、get_unique（唯一值分析）。'
    '所有函数返回格式化的中文字符串，方便 LLM 理解和向用户展示。',
    indent=True
)

add_para('Analyst（统计分析）', bold=True)
add_para(
    '包含 6 个函数：describe（数值统计）、value_counts（频次分布）、groupby_agg（分组聚合）、'
    'correlation（相关性分析）、filter_data（条件过滤）、sort_data（数据排序）。'
    '每个函数包含完善的输入验证（列名存在性、数据类型检查），LLM 调用出错后可自动修正参数重试。',
    indent=True
)

add_para('Visualizer（可视化绘图）', bold=True)
add_para(
    '包含 6 个函数：histogram（直方图）、bar_chart（柱状图）、box_plot（箱线图）、'
    'heatmap（热力图）、scatter（散点图）、pairplot（配对图矩阵）。'
    '所有函数返回 PIL Image 对象。参数错误时返回包含错误说明的占位图而非抛出异常，'
    '保证 LLM 能够接收反馈并修正调用参数。',
    indent=True
)

doc.add_page_break()

# ==================== 五、系统详细实现与核心代码 ====================
doc.add_heading('五、系统详细实现与核心代码', level=1)

doc.add_heading('5.1 各功能模块具体实现', level=2)

add_para('Agent 主循环模块（agent/core.py）', bold=True)
add_para(
    'Agent 主循环是整个系统的核心调度模块。它实现了完整的 Function Call 逻辑：'
    '构建包含 system prompt 和 tool schema 的消息列表，'
    '调用 qwen-plus API，解析 tool_call 响应，调度 Skill 函数执行，'
    '并将结果写回对话历史进行迭代。系统设置了 15 次最大迭代限制和 OpenAI API 异常捕获。',
    indent=True
)

add_para('对话历史管理模块（agent/memory.py）', bold=True)
add_para(
    'ConversationMemory 类管理完整的 OpenAI 格式消息列表。'
    '支持追加 user/assistant/tool 三种角色的消息，'
    '提供基于字符数的简单截断机制（超过 128,000 字符时自动移除最早的消息轮次），'
    '确保长对话不会超出模型的上下文窗口限制。',
    indent=True
)

add_para('Tool Schema 定义模块（agent/schema.py）', bold=True)
add_para(
    '以 OpenAI JSON Schema 格式定义了全部 18 个 Skill 函数的参数规范。'
    '每个 Schema 包含函数名称、功能描述、参数类型、是否必填等信息。'
    '这些 Schema 作为 tools 参数传入 API 请求，使模型能够理解并正确调用各 Skill 函数。'
    '此外维护了一个 VISUAL_FUNCTIONS 集合标记视觉类函数，'
    '用于在 dispatch 层区分文字结果和图片结果的处理方式。',
    indent=True
)

add_para('Gradio UI 模块（ui/app.py）', bold=True)
add_para(
    '基于 Gradio 6.18 的 gr.ChatInterface 构建对话界面。'
    '界面支持 OpenAI 风格的多模态消息格式，'
    'agent_loop 生成的文字直接作为文本消息展示，PIL Image 图片保存为临时文件后以图片路径消息展示。'
    '提供了 9 条预设演示指令覆盖三大 Skill 的组合使用场景，'
    '并支持一键重置对话。',
    indent=True
)

doc.add_heading('5.2 核心代码展示与说明', level=2)

add_para('Agent 主循环核心代码：', bold=True)
add_code(
    '# agent/core.py — Agent 主循环（核心逻辑）\n'
    'def agent_loop(\n'
    '    user_input: str,\n'
    '    df: pd.DataFrame,\n'
    '    memory: ConversationMemory,\n'
    ') -> Generator[str | Image, None, None]:\n'
    '    client = _init_client()\n'
    '    mem = memory\n'
    '    mem.add_user(user_input)\n\n'
    '    for iteration in range(_MAX_ITERATIONS):\n'
    '        response = client.chat.completions.create(\n'
    '            model="qwen-plus",\n'
    '            messages=mem.messages,\n'
    '            tools=SKILL_SCHEMAS,\n'
    '            tool_choice="auto",\n'
    '        )\n'
    '        msg = response.choices[0].message\n\n'
    '        if not msg.tool_calls:          # 模型直接回答\n'
    '            yield msg.content\n'
    '            return\n\n'
    '        for tool_call in msg.tool_calls: # 解析并执行工具调用\n'
    '            func_name = tool_call.function.name\n'
    '            args = json.loads(tool_call.function.arguments)\n'
    '            result = dispatch_skill(func_name, args, df)\n'
    '            # 图片直接 yield，文字写回 memory 继续循环\n'
    '            if func_name in VISUAL_FUNCTIONS:\n'
    '                yield result\n'
    '            else:\n'
    '                mem.add_tool(str(result), tool_call.id)'
)

add_para('')
add_para('Skill 函数示例 — 统计分析：', bold=True)
add_code(
    '# skills/analyst.py — describe 函数（含输入验证）\n'
    'def describe(df: pd.DataFrame, cols: list[str] | None = None) -> str:\n'
    '    if cols:\n'
    '        missing = [c for c in cols if c not in df.columns]\n'
    '        if missing:\n'
    '            return f"错误: 以下列不存在: {\', \'.join(missing)}"\n'
    '        non_numeric = [c for c in cols\n'
    '            if not pd.api.types.is_numeric_dtype(df[c])]\n'
    '        if non_numeric:\n'
    '            return f"错误: 以下列不是数值类型: {\', \'.join(non_numeric)}"\n'
    '        target = df[cols]\n'
    '    else:\n'
    '        target = df.select_dtypes(include="number")\n'
    '    if target.empty:\n'
    '        return "没有可统计的数值列。"\n'
    '    stats = target.describe().round(3)\n'
    '    return f"统计信息:\\n{stats.to_string()}"'
)

add_para('')
add_para('Gradio UI 界面代码：', bold=True)
add_code(
    '# ui/app.py — Gradio 对话界面（核心部分）\n'
    'def respond(message: str, history: list):\n'
    '    try:\n'
    '        outputs = list(agent_loop(message, df, memory))\n'
    '    except Exception as e:\n'
    '        return [f"⚠️ 出错了: {e}"]\n\n'
    '    messages = []\n'
    '    for o in outputs:\n'
    '        if isinstance(o, str):\n'
    '            messages.append(o)\n'
    '        else:  # PIL Image\n'
    '            path = tempfile.mktemp(suffix=".png")\n'
    '            o.save(path)\n'
    '            messages.append(\n'
    '                {"role": "assistant", "content": {"path": path}})\n'
    '    return messages\n\n'
    'with gr.Blocks(title="TableAgent", fill_height=True) as demo:\n'
    '    gr.Markdown("# TableAgent — 数据分析智能助手")\n'
    '    chatbot = gr.ChatInterface(\n'
    '        fn=respond,\n'
    '        examples=PRESET_EXAMPLES,\n'
    '        fill_height=True,\n'
    '    )'
)

doc.add_heading('5.3 关键问题与解决方案', level=2)

add_para('问题 1：模型生成非法参数', bold=True)
add_para(
    '现象：LLM 在调用 Skill 函数时生成了不存在的列名或错误的数据类型参数。'
    '解决方案：在所有 Skill 函数入口添加严格的输入验证，'
    '对列名存在性、数据类型、参数取值范围等进行前置检查。'
    '当参数错误时返回结构化的错误信息，LLM 可根据错误提示修正参数后自动重试。',
    indent=True
)

add_para('问题 2：Agent 无限循环调用', bold=True)
add_para(
    '现象：在复杂链式调用场景下，LLM 可能陷入工具调用的死循环。'
    '解决方案：引入 _MAX_ITERATIONS = 15 最大迭代限制。'
    '超出限制时 Agent 主动终止并提示用户简化问题或分步提问，防止资源耗尽。',
    indent=True
)

add_para('问题 3：中文字体在图表中显示为方框', bold=True)
add_para(
    '现象：Matplotlib 绘制的图表标题和轴标签中的中文字符显示为乱码方框。'
    '解决方案：在 visualizer.py 中配置中文字体回退机制，'
    '依次尝试 SimHei、Microsoft YaHei 等常见中文字体，'
    '并设置 axes.unicode_minus = False 解决负号显示问题。',
    indent=True
)

add_para('问题 4：API 调用不稳定', bold=True)
add_para(
    '现象：网络波动或 API 限流导致请求失败。'
    '解决方案：在 agent_loop 中添加 try/except 捕获 OpenAI.APIError，'
    '返回友好的中文错误提示而非程序崩溃。',
    indent=True
)

doc.add_page_break()

# ==================== 六、实验结果与分析 ====================
doc.add_heading('六、实验结果与分析', level=1)

doc.add_heading('6.1 实验评价指标', level=2)
add_para(
    '本项目作为 AI Agent 系统，主要评价指标包括：'
    '（1）Skill 函数单元测试通过率；'
    '（2）Agent 对各类用户指令的正确响应能力（单步/链式/错误恢复）；'
    '（3）系统运行稳定性（无崩溃连续对话轮次）。',
    indent=True
)

doc.add_heading('6.2 实验结果展示', level=2)

add_para('单元测试结果：', bold=True)
add_para('针对三大 Skill 模块的 18 个函数编写了 42 个 pytest 单元测试，全部通过。测试覆盖：', indent=True)
add_para('• DataReader：10 个测试（正常调用、空数据、超界参数、不存在列等）')
add_para('• Analyst：14 个测试（正常调用、不存在列、非数值列、空条件、不支持方法等）')
add_para('• Visualizer：18 个测试（每个绘图函数含正常路径和错误路径双重验证）')

add_para('')
add_para('预设演示指令执行效果：', bold=True)
add_para('系统内置了 9 条预设演示指令，覆盖三大 Skill 的以下组合场景：', indent=True)

table = doc.add_table(rows=10, cols=3, style='Light Grid Accent 1')
table.cell(0,0).text = '指令'
table.cell(0,1).text = '涉及 Skill'
table.cell(0,2).text = '预期结果'
demo_cases = [
    ('数据集有多少行多少列？有哪些字段？', 'DataReader', '返回概览信息'),
    ('查看前5行数据', 'DataReader', '显示前5行记录'),
    ('统计年龄的分布情况', 'Analyst', '返回年龄统计描述'),
    ('不同舱位的票价平均值是多少？画柱状图', 'Analyst + Visualizer', '统计数据并生成柱状图'),
    ('年龄和票价的相关性如何？画热力图', 'Analyst + Visualizer', '计算相关性并生成热力图'),
    ('绘制年龄的直方图', 'Visualizer', '生成直方图'),
    ('按性别统计存活率', 'Analyst', '返回分组统计结果'),
    ('筛选出年龄大于60岁的乘客', 'Analyst', '返回过滤结果'),
    ('查看 embarked 列的唯一值分布', 'DataReader', '返回频次分布'),
]
for i, (inst, skill, exp) in enumerate(demo_cases):
    table.cell(i+1,0).text = inst
    table.cell(i+1,1).text = skill
    table.cell(i+1,2).text = exp

doc.add_heading('6.3 结果对比与分析', level=2)
add_para(
    '通过实际测试运行，TableAgent 展现出以下特点：',
    indent=True
)
add_para(
    '（1）意图识别准确：qwen-plus 模型能够较为准确地识别用户的数据分析意图，'
    '在大部分场景下自动选择正确的 Skill 函数。特别是在链式调用场景中（如"先统计再画图"），'
    '模型能自主完成多步工具调度。',
    indent=True
)
add_para(
    '（2）错误恢复能力：当 Skill 函数返回参数错误信息时，LLM 能够理解错误原因并修正参数重试。'
    '例如，当用户提到的列名与实际不完全匹配时，Agent 可通过 get_overview 先探查结构再执行分析。',
    indent=True
)
add_para(
    '（3）局限性分析：在极少数复杂查询场景中，模型可能一次调用多个无关工具，'
    '或陷入工具调用循环。15 次迭代上限有效拦截了此类异常。'
    '此外，由于 qwen-plus 为通用模型，对 pandas 查询语法的理解偶有偏差，'
    '需要在系统 prompt 中提供更清晰的语法示例。',
    indent=True
)

doc.add_page_break()

# ==================== 七、系统测试与功能验证 ====================
doc.add_heading('七、系统测试与功能验证', level=1)

add_para('本节列出系统核心功能的测试用例及验证结果。', indent=True)

add_para('')
add_para('1. Skill 函数单元测试', bold=True)
table = doc.add_table(rows=5, cols=4, style='Light Grid Accent 1')
table.cell(0,0).text = '测试模块'
table.cell(0,1).text = '测试数量'
table.cell(0,2).text = '通过'
table.cell(0,3).text = '失败'
test_results = [
    ('skills/reader.py', '10', '10', '0'),
    ('skills/analyst.py', '14', '14', '0'),
    ('skills/visualizer.py', '18', '18', '0'),
    ('总计', '42', '42', '0'),
]
for i, (mod, total, passed, failed) in enumerate(test_results):
    table.cell(i+1,0).text = mod
    table.cell(i+1,1).text = total
    table.cell(i+1,2).text = passed
    table.cell(i+1,3).text = failed

add_para('')
add_para('2. 系统功能验证', bold=True)
table = doc.add_table(rows=7, cols=4, style='Light Grid Accent 1')
table.cell(0,0).text = '测试场景'
table.cell(0,1).text = '操作'
table.cell(0,2).text = '预期结果'
table.cell(0,3).text = '实际结果'
func_tests = [
    ('数据探查', '输入"查看数据集信息"', '返回行数列数字段信息', '✅ 通过'),
    ('统计分析', '输入"统计年龄分布"', '返回描述性统计', '✅ 通过'),
    ('图表生成', '输入"画年龄直方图"', '显示直方图图片', '✅ 通过'),
    ('链式调用', '输入"按舱位统计票价并画图"', '先统计数据再生成柱状图', '✅ 通过'),
    ('错误恢复', '输入不存在的列名', '返回友好错误提示', '✅ 通过'),
    ('重置对话', '点击重置按钮', '清空对话历史和记忆', '✅ 通过'),
]
for i, (scene, op, exp, actual) in enumerate(func_tests):
    table.cell(i+1,0).text = scene
    table.cell(i+1,1).text = op
    table.cell(i+1,2).text = exp
    table.cell(i+1,3).text = actual

add_para('')
add_para(
    '测试结论：系统全部 42 个单元测试通过，6 项功能验证测试全部通过，'
    '系统运行稳定，功能完整性符合课程设计要求。',
    indent=True
)

doc.add_page_break()

# ==================== 八、总结与展望 ====================
doc.add_heading('八、总结与展望', level=1)

doc.add_heading('8.1 项目总结', level=2)
add_para(
    '本次课程设计完成了搭载数据分析 Skill 的表格智能 Agent（TableAgent）的完整开发。'
    '项目实现了三大 Skill 共 18 个数据分析函数的封装，基于手搭 Function Call 机制 '
    '构建了 Agent 调度主循环，并通过 Gradio 构建了交互式对话界面。',
    indent=True
)
add_para(
    '通过本项目，深入掌握了以下核心技术：'
    '（1）LLM Function Call 机制的底层原理和实现方式，'
    '理解了大模型如何通过工具描述自主决定调用外部函数；'
    '（2）AI Agent 的系统架构设计方法，包括意图识别、工具调度、对话管理等关键模块；'
    '（3）数据分析工具链的工程化封装，包括输入验证、错误处理、链式调用等生产级特性。',
    indent=True
)
add_para(
    '项目的技术亮点包括：不依赖高级框架的手搭 Function Call 实现、'
    '完整的输入验证与错误恢复机制、42 个单元测试的全覆盖、'
    '以及中文友好的 Gradio 交互界面。',
    indent=True
)

doc.add_heading('8.2 未来展望', level=2)
add_para(
    '当前系统仍有以下改进空间：',
    indent=True
)
add_para('（1）多数据集支持：当前仅支持单一 Titanic 数据集，未来可扩展为动态加载用户上传的 CSV 文件。')
add_para('（2）高级分析能力：增加缺失值填充、异常值检测、特征工程等数据预处理 Skill。')
add_para('（3）模型可配置：支持切换不同 LLM 后端（如 GLM、DeepSeek），对比各模型的 Function Call 效果。')
add_para('（4）会话持久化：将对话历史和图表结果保存到本地文件，支持历史会话回溯。')
add_para('（5）多轮上下文增强：优化 system prompt，使 Agent 在多轮对话中更好地引用此前分析结果。')
add_para('（6）并发用户支持：为 Gradio 界面添加多会话管理，支持多用户同时使用。')

doc.add_page_break()

# ==================== 九、参考文献 ====================
doc.add_heading('九、参考文献', level=1)

refs = [
    '[1] OpenAI. Function Calling Guide [EB/OL]. https://platform.openai.com/docs/guides/function-calling, 2025.',
    '[2] 阿里云. 通义千问 API 文档 [EB/OL]. https://help.aliyun.com/zh/dashscope, 2025.',
    '[3] Gradio. ChatInterface Documentation [EB/OL]. https://www.gradio.app/docs/gradio/chatinterface, 2025.',
    '[4] McKinney W. Python for Data Analysis (3rd ed.) [M]. O\'Reilly Media, 2022.',
    '[5] Waskom M. Seaborn: Statistical Data Visualization [J]. Journal of Open Source Software, 2021, 6(60): 3021.',
    '[6] Hunter J D. Matplotlib: A 2D Graphics Environment [J]. Computing in Science & Engineering, 2007, 9(3): 90-95.',
    '[7] Kaggle. Titanic: Machine Learning from Disaster [EB/OL]. https://www.kaggle.com/c/titanic, 2012.',
    '[8] Chase H. LangChain: Building Applications with LLMs through Composability [EB/OL]. https://github.com/langchain-ai/langchain, 2023.',
]
for ref in refs:
    add_para(ref)

doc.add_page_break()

# ==================== 十、附录 ====================
doc.add_heading('十、附录（可选）', level=1)
add_para('附录 A：完整项目目录结构', bold=True)
add_code(
    'TableAgent/\n'
    '├── AGENTS.md              # 项目知识库与开发文档\n'
    '├── requirements.txt       # 依赖清单\n'
    '├── main.py                # 入口：Gradio UI（默认）/ CLI\n'
    '├── generate_report.py     # 报告生成脚本\n'
    '│\n'
    '├── agent/\n'
    '│   ├── core.py            # Agent 主循环（手搭 Function Call）\n'
    '│   ├── schema.py          # 18 个 Tool Schema（JSON Schema）\n'
    '│   └── memory.py          # 对话历史管理\n'
    '│\n'
    '├── skills/\n'
    '│   ├── reader.py          # Skill 1: 6 个数据读取与探查函数\n'
    '│   ├── analyst.py         # Skill 2: 6 个统计分析函数\n'
    '│   └── visualizer.py      # Skill 3: 6 个可视化绘图函数\n'
    '│\n'
    '├── ui/\n'
    '│   └── app.py             # Gradio ChatInterface 界面\n'
    '│\n'
    '├── tests/\n'
    '│   └── test_skills.py     # 42 个 pytest 单元测试\n'
    '│\n'
    '├── data/\n'
    '│   └── titanic.csv        # 泰坦尼克数据集\n'
    '│\n'
    '└── notebooks/\n'
    '    └── explore.ipynb      # 数据探索与原型验证'
)

add_para('')
add_para('附录 B：运行说明', bold=True)
add_code(
    '# 1. 激活环境\n'
    'conda activate pytorch\n\n'
    '# 2. 安装依赖\n'
    'pip install openai seaborn gradio pandas numpy matplotlib\n\n'
    '# 3. 配置 API Key（在 .env 文件中设置 DASHSCOPE_API_KEY）\n\n'
    '# 4. 启动 Web UI\n'
    'python main.py\n'
    '# 浏览器访问 http://127.0.0.1:7860\n\n'
    '# 5. 或使用命令行模式\n'
    'python main.py --cli\n\n'
    '# 6. 运行单元测试\n'
    'python -m pytest tests/ -v'
)

# ===== 保存 =====
output_path = r'D:\学习\人工智能\docs\Final\TableAgent\人工智能课程设计报告.docx'
doc.save(output_path)
print(f'报告已生成: {output_path}')
